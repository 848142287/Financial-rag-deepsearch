"""
Word文档解析器
处理.doc和.docx格式的文档
"""

import logging
from typing import Dict, List, Any, Optional
from .base import BaseFileParser, ParseResult

logger = logging.getLogger(__name__)

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available, .docx parsing will be limited")

try:
    import win32com.client as win32
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False
    logger.warning("pywin32 not available, .doc parsing will not be available")


class DocParser(BaseFileParser):
    """Word文档解析器"""

    def __init__(self):
        self.supported_extensions = ['.doc', '.docx']
        self.supported_mime_types = [
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]

    def parse(self, file_path: str, **kwargs) -> ParseResult:
        """
        解析Word文档

        Args:
            file_path: 文件路径
            **kwargs: 其他参数

        Returns:
            ParseResult: 解析结果
        """
        try:
            import os
            _, ext = os.path.splitext(file_path.lower())

            if ext == '.docx':
                return self._parse_docx(file_path, **kwargs)
            elif ext == '.doc':
                return self._parse_doc(file_path, **kwargs)
            else:
                return ParseResult(
                    content="",
                    metadata={'file_type': 'doc', 'error': f'不支持的文件扩展名: {ext}'},
                    success=False,
                    error_message=f'不支持的文件扩展名: {ext}'
                )

        except Exception as e:
            logger.error(f"Word文档解析失败: {e}")
            return ParseResult(
                content="",
                metadata={'file_type': 'doc', 'error': str(e)},
                success=False,
                error_message=str(e)
            )

    def _parse_docx(self, file_path: str, **kwargs) -> ParseResult:
        """解析.docx文件"""
        if not DOCX_AVAILABLE:
            return ParseResult(
                content="",
                metadata={'file_type': 'docx', 'error': 'python-docx not available'},
                success=False,
                error_message='python-docx library not installed'
            )

        try:
            doc = Document(file_path)

            # 提取文本内容
            content_parts = []

            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    content_parts.append(para.text)

            # 提取表格内容
            table_count = 0
            for table in doc.tables:
                table_count += 1
                content_parts.append(f"\n表格 {table_count}:")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content_parts.append(" | ".join(row_text))
                content_parts.append("")  # 空行分隔表格

            content = "\n".join(content_parts)

            # 构建元数据
            metadata = {
                'file_type': 'docx',
                'paragraphs': len(doc.paragraphs),
                'tables': table_count,
                'sections': len(doc.sections),
                'core_properties': self._get_core_properties(doc)
            }

            return ParseResult(
                content=content,
                metadata=metadata,
                success=True
            )

        except Exception as e:
            logger.error(f"DOCX解析失败: {e}")
            return ParseResult(
                content="",
                metadata={'file_type': 'docx', 'error': str(e)},
                success=False,
                error_message=str(e)
            )

    def _parse_doc(self, file_path: str, **kwargs) -> ParseResult:
        """解析.doc文件（需要Windows环境）"""
        if not WIN32_AVAILABLE:
            # 如果没有win32，尝试用其他方法或返回基本信息
            return self._parse_doc_basic(file_path)

        try:
            # 使用COM接口解析.doc文件
            word = win32.Dispatch("Word.Application")
            word.Visible = False

            doc = word.Documents.Open(file_path)
            content = doc.Content.Text

            doc.Close()
            word.Quit()

            metadata = {
                'file_type': 'doc',
                'pages': 0,  # 无法直接获取页数
                'words': len(content.split()),
                'characters': len(content)
            }

            return ParseResult(
                content=content,
                metadata=metadata,
                success=True
            )

        except Exception as e:
            logger.error(f"DOC解析失败: {e}")
            return self._parse_doc_basic(file_path)

    def _parse_doc_basic(self, file_path: str) -> ParseResult:
        """基础.doc文件解析（当COM不可用时）"""
        import os

        file_size = os.path.getsize(file_path)

        content = f"Word文档: {os.path.basename(file_path)}\n"
        content += f"文件大小: {file_size} 字节\n"
        content += "注意: 由于缺少pywin32库，无法提取.doc文件内容，请转换为.docx格式"

        metadata = {
            'file_type': 'doc',
            'file_size': file_size,
            'parsing_limited': True,
            'recommendation': 'Convert to .docx format for full parsing'
        }

        return ParseResult(
            content=content,
            metadata=metadata,
            success=True
        )

    def _get_core_properties(self, doc) -> Dict[str, Any]:
        """获取文档核心属性"""
        try:
            props = doc.core_properties
            return {
                'title': getattr(props, 'title', '') or '',
                'author': getattr(props, 'author', '') or '',
                'subject': getattr(props, 'subject', '') or '',
                'created': str(getattr(props, 'created', '')) or '',
                'modified': str(getattr(props, 'modified', '')) or '',
                'keywords': getattr(props, 'keywords', '') or ''
            }
        except Exception as e:
            logger.warning(f"无法获取文档属性: {e}")
            return {}

    def is_supported(self, file_path: str, mime_type: str = None) -> bool:
        """
        检查是否支持解析该文件

        Args:
            file_path: 文件路径
            mime_type: MIME类型（可选）

        Returns:
            bool: 是否支持
        """
        import os
        _, ext = os.path.splitext(file_path.lower())
        if ext in self.supported_extensions:
            return True

        if mime_type and mime_type in self.supported_mime_types:
            return True

        return False