"""
多格式文档加载器 - Multi-Format Document Loader

支持多种文档格式的统一加载接口：
- PDF (使用pymupdf4llm)
- DOCX/DOC (使用python-docx)
- Markdown
- 纯文本

所有文档统一转换为Markdown格式输出，便于后续处理
"""

from abc import ABC, abstractmethod
from pathlib import Path
from dataclasses import dataclass

from app.core.structured_logging import get_structured_logger

logger = get_structured_logger(__name__)

@dataclass
class DocumentLoadResult:
    """文档加载结果"""
    content: str  # Markdown格式内容
    metadata: Dict[str, Any]
    num_pages: int = 0
    file_size: int = 0

class BaseLoader(ABC):
    """文档加载器基类"""

    @abstractmethod
    def load(self, file_path: str) -> DocumentLoadResult:
        """
        加载文档

        Args:
            file_path: 文件路径

        Returns:
            DocumentLoadResult
        """
        pass

    @property
    @abstractmethod
    def supported_extensions(self) -> List[str]:
        """支持的文件扩展名列表（包含点号）"""
        pass

    def _get_file_metadata(self, file_path: str) -> Dict[str, Any]:
        """获取文件元数据"""
        path = Path(file_path)
        return {
            "filename": path.name,
            "file_type": path.suffix.lower(),
            "file_size": path.stat().st_size if path.exists() else 0,
        }

class PDFLoader(BaseLoader):
    """PDF文档加载器（使用pymupdf4llm）"""

    @property
    def supported_extensions(self) -> List[str]:
        return [".pdf"]

    def load(self, file_path: str) -> DocumentLoadResult:
        """使用pymupdf4llm将PDF转换为Markdown"""
        try:
            import pymupdf4llm
        except ImportError:
            raise ImportError(
                "pymupdf4llm未安装，请运行: pip install pymupdf4llm"
            )

        try:
            # 转换为Markdown
            md_text = pymupdf4llm.to_markdown(file_path)

            # 获取页数（使用pymupdf4llm的底层fitz）
            import fitz
            doc = fitz.open(file_path)
            num_pages = doc.page_count
            doc.close()

            metadata = self._get_file_metadata(file_path)
            metadata["num_pages"] = num_pages

            logger.info(f"✅ PDF加载成功: {file_path} ({num_pages}页)")

            return DocumentLoadResult(
                content=md_text,
                metadata=metadata,
                num_pages=num_pages,
                file_size=metadata["file_size"]
            )

        except Exception as e:
            logger.error(f"PDF加载失败: {e}")
            raise ValueError(f"无法加载PDF文件: {e}")

class DocxLoader(BaseLoader):
    """DOCX/DOC文档加载器（使用python-docx）"""

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    def load(self, file_path: str) -> DocumentLoadResult:
        """使用python-docx加载DOCX"""
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx未安装，请运行: pip install python-docx"
            )

        try:
            doc = docx.Document(file_path)

            paragraphs = []
            num_paragraphs = len(doc.paragraphs)

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 识别标题样式并转换为Markdown标题
                    style_name = para.style.name

                    if style_name.startswith('Heading'):
                        # 提取标题级别（如 "Heading 1" -> 1）
                        try:
                            level = int(style_name.split()[-1])
                        except (ValueError, IndexError):
                            level = 1

                        # 转换为Markdown标题
                        paragraphs.append(f"{'#' * level} {text}")

                    elif style_name == 'Title':
                        paragraphs.append(f"# {text}")

                    else:
                        # 普通段落
                        paragraphs.append(text)

            # 连接所有段落
            content = "\n\n".join(paragraphs)

            metadata = self._get_file_metadata(file_path)
            metadata["num_paragraphs"] = num_paragraphs

            logger.info(f"✅ DOCX加载成功: {file_path} ({num_paragraphs}段)")

            return DocumentLoadResult(
                content=content,
                metadata=metadata,
                file_size=metadata["file_size"]
            )

        except Exception as e:
            logger.error(f"DOCX加载失败: {e}")
            raise ValueError(f"无法加载DOCX文件: {e}")

class MarkdownLoader(BaseLoader):
    """Markdown文档加载器"""

    @property
    def supported_extensions(self) -> List[str]:
        return [".md", ".markdown"]

    def load(self, file_path: str) -> DocumentLoadResult:
        """直接加载Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 统计基本信息
            num_lines = len(content.split('\n'))
            num_headings = content.count('\n#')

            metadata = self._get_file_metadata(file_path)
            metadata["num_lines"] = num_lines
            metadata["num_headings"] = num_headings

            logger.info(f"✅ Markdown加载成功: {file_path}")

            return DocumentLoadResult(
                content=content,
                metadata=metadata,
                file_size=metadata["file_size"]
            )

        except Exception as e:
            logger.error(f"Markdown加载失败: {e}")
            raise ValueError(f"无法加载Markdown文件: {e}")

class TextLoader(BaseLoader):
    """纯文本文档加载器"""

    @property
    def supported_extensions(self) -> List[str]:
        return [".txt", ".text"]

    def load(self, file_path: str) -> DocumentLoadResult:
        """加载纯文本文件"""
        try:
            # 尝试多种编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']

            content = None
            used_encoding = None

            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue

            if content is None:
                raise ValueError("无法用常见编码解码文件")

            num_lines = len(content.split('\n'))
            num_chars = len(content)

            metadata = self._get_file_metadata(file_path)
            metadata["encoding"] = used_encoding
            metadata["num_lines"] = num_lines
            metadata["num_chars"] = num_chars

            logger.info(f"✅ 文本加载成功: {file_path} (编码: {used_encoding})")

            return DocumentLoadResult(
                content=content,
                metadata=metadata,
                file_size=metadata["file_size"]
            )

        except Exception as e:
            logger.error(f"文本加载失败: {e}")
            raise ValueError(f"无法加载文本文件: {e}")

class MultiFormatDocumentLoader:
    """
    多格式文档加载器

    统一接口加载多种格式的文档，自动识别文件类型并选择合适的加载器
    """

    def __init__(self):
        """初始化加载器"""
        self.loaders = {
            ".pdf": PDFLoader(),
            ".docx": DocxLoader(),
            ".doc": DocxLoader(),
            ".md": MarkdownLoader(),
            ".markdown": MarkdownLoader(),
            ".txt": TextLoader(),
            ".text": TextLoader(),
        }

        logger.info(f"多格式文档加载器初始化完成，支持格式: {list(self.loaders.keys())}")

    def load(self, file_path: str) -> DocumentLoadResult:
        """
        自动识别文件格式并加载

        Args:
            file_path: 文件路径

        Returns:
            DocumentLoadResult

        Raises:
            ValueError: 不支持的文件格式
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = path.suffix.lower()

        if ext not in self.loaders:
            raise ValueError(
                f"不支持的文件格式: {ext}。"
                f"支持的格式: {list(self.loaders.keys())}"
            )

        loader = self.loaders[ext]
        return loader.load(file_path)

    def load_batch(
        self,
        file_paths: List[str],
        ignore_errors: bool = True
    ) -> List[DocumentLoadResult]:
        """
        批量加载文档

        Args:
            file_paths: 文件路径列表
            ignore_errors: 是否忽略错误（True: 跳过失败文件，False: 遇到错误抛出异常）

        Returns:
            DocumentLoadResult列表
        """
        results = []
        errors = []

        for file_path in file_paths:
            try:
                result = self.load(file_path)
                results.append(result)
            except Exception as e:
                if ignore_errors:
                    logger.warning(f"跳过文件 {file_path}: {e}")
                    errors.append({"file": file_path, "error": str(e)})
                else:
                    raise

        logger.info(f"✅ 批量加载完成: 成功 {len(results)}/{len(file_paths)}")

        if errors and ignore_errors:
            logger.warning(f"⚠️ {len(errors)} 个文件加载失败")

        return results

    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名列表"""
        return list(self.loaders.keys())

    def is_supported(self, file_path: str) -> bool:
        """检查文件是否支持"""
        ext = Path(file_path).suffix.lower()
        return ext in self.loaders

# 便捷函数
def load_document(file_path: str) -> DocumentLoadResult:
    """
    加载文档的便捷函数

    Args:
        file_path: 文件路径

    Returns:
        DocumentLoadResult
    """
    loader = MultiFormatDocumentLoader()
    return loader.load(file_path)

def load_documents_batch(
    file_paths: List[str],
    ignore_errors: bool = True
) -> List[DocumentLoadResult]:
    """
    批量加载文档的便捷函数

    Args:
        file_paths: 文件路径列表
        ignore_errors: 是否忽略错误

    Returns:
        DocumentLoadResult列表
    """
    loader = MultiFormatDocumentLoader()
    return loader.load_batch(file_paths, ignore_errors=ignore_errors)
