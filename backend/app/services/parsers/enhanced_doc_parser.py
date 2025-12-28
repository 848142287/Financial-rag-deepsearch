"""
增强Word文档解析器 - 支持分章节、多模态分析、结构化输出
"""

import logging
import os
import tempfile
import shutil
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from dataclasses import dataclass, field
from datetime import datetime
import asyncio
import re

from .base import BaseFileParser, ParseResult
from .multimodal_analyzer import get_multimodal_analyzer, ImageAnalysisResult

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")


@dataclass
class DocumentSection:
    """文档章节"""
    level: int  # 标题级别 1-6
    title: str  # 章节标题
    content: List[str]  # 章节内容（文本段落）
    subsections: List['DocumentSection'] = field(default_factory=list)  # 子章节
    images: List[Dict[str, Any]] = field(default_factory=list)  # 图片
    tables: List[Dict[str, Any]] = field(default_factory=list)  # 表格
    formulas: List[Dict[str, Any]] = field(default_factory=list)  # 公式
    metadata: Dict[str, Any] = field(default_factory=dict)  # 元数据


@dataclass
class DocumentElement:
    """文档元素"""
    element_type: str  # 'heading', 'paragraph', 'image', 'table', 'formula', 'chart'
    content: str  # 内容或描述
    level: int = 0  # 标题级别
    metadata: Dict[str, Any] = field(default_factory=dict)  # 元数据
    position: int = 0  # 在文档中的位置


class EnhancedDocParser(BaseFileParser):
    """增强Word文档解析器"""

    def __init__(self, config: Optional[Dict[str, Any]] = None):
        super().__init__(config)
        self.enable_multimodal = self.config.get('enable_multimodal', True)
        self.extract_images = self.config.get('extract_images', True)
        self.enable_ocr = self.config.get('enable_ocr', True)

        # 初始化多模态分析器
        if self.enable_multimodal:
            self.multimodal_analyzer = get_multimodal_analyzer(self.config)
        else:
            self.multimodal_analyzer = None

    @property
    def supported_extensions(self) -> List[str]:
        """支持的文件扩展名列表"""
        return ['.docx']

    @property
    def parser_name(self) -> str:
        """解析器名称"""
        return "EnhancedDocParser"

    def can_parse(self, file_path: str, file_extension: str = None) -> bool:
        """检查是否能解析指定文件"""
        if not DOCX_AVAILABLE:
            return False

        if file_extension:
            return file_extension.lower() in self.supported_extensions

        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.supported_extensions

    async def parse(self, file_path: str, **kwargs) -> ParseResult:
        """
        解析Word文档

        Args:
            file_path: Word文件路径
            **kwargs: 其他参数
                - extract_images: 是否提取图像（默认True）
                - enable_ocr: 是否启用OCR（默认True）
                - output_format: 输出格式（'markdown', 'json'，默认'markdown'）

        Returns:
            ParseResult: 解析结果
        """
        import time
        start_time = time.time()

        try:
            if not DOCX_AVAILABLE:
                return ParseResult(
                    content="",
                    metadata={'file_type': 'docx', 'error': 'python-docx library not available'},
                    success=False,
                    error_message='python-docx library not installed'
                )

            # 获取参数
            extract_images = kwargs.get('extract_images', self.extract_images)
            enable_ocr = kwargs.get('enable_ocr', self.enable_ocr)
            output_format = kwargs.get('output_format', 'markdown')

            # 创建临时目录用于提取图片
            temp_dir = None
            if extract_images:
                temp_dir = tempfile.mkdtemp(prefix='docx_images_')

            # 1. 解析文档结构
            elements = await self._parse_document_structure(file_path, temp_dir)

            # 2. 组织为章节结构
            sections = self._organize_sections(elements)

            # 3. 分析图片、图表、公式
            if extract_images and self.multimodal_analyzer:
                await self._analyze_multimedia_elements(elements, temp_dir, enable_ocr)

            # 4. 生成输出
            if output_format == 'markdown':
                content = self._generate_markdown(elements, sections)
            else:
                content = self._generate_json_output(elements, sections)

            # 5. 清理临时目录
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)

            # 6. 构建元数据
            metadata = self._build_metadata(file_path, elements, sections)
            metadata['parse_time'] = time.time() - start_time

            return ParseResult(
                content=content,
                metadata=metadata,
                success=True,
                parse_time=time.time() - start_time
            )

        except Exception as e:
            logger.error(f"Enhanced Word parsing failed: {e}", exc_info=True)

            # 清理临时目录
            if temp_dir and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir, ignore_errors=True)

            return ParseResult(
                content="",
                metadata={'file_type': 'docx', 'error': str(e)},
                success=False,
                error_message=str(e),
                parse_time=time.time() - start_time
            )

    async def _parse_document_structure(
        self,
        file_path: str,
        temp_dir: Optional[str]
    ) -> List[DocumentElement]:
        """解析文档结构"""
        doc = Document(file_path)
        elements = []
        position = 0

        # 遍历文档的所有元素
        for element in doc.element.body:
            if isinstance(element, CT_P):
                # 处理段落
                paragraph_elements = self._process_paragraph(element, doc, temp_dir, position)
                elements.extend(paragraph_elements)
                position += len(paragraph_elements)

            elif isinstance(element, CT_Tbl):
                # 处理表格
                table_element = self._process_table(element, doc, position)
                if table_element:
                    elements.append(table_element)
                    position += 1

        return elements

    def _process_paragraph(
        self,
        ct_p: CT_P,
        doc: Document,
        temp_dir: Optional[str],
        position: int
    ) -> List[DocumentElement]:
        """处理段落"""
        elements = []

        try:
            paragraph = Paragraph(ct_p, doc)

            # 获取段落样式
            style_name = paragraph.style.name if paragraph.style else ""

            # 判断是否为标题
            if 'Heading' in style_name:
                level = self._extract_heading_level(style_name)
                elements.append(DocumentElement(
                    element_type='heading',
                    content=paragraph.text,
                    level=level,
                    metadata={'style': style_name},
                    position=position
                ))
            elif paragraph.text.strip():
                # 普通段落
                # 检查是否包含图片
                images = self._extract_images_from_paragraph(paragraph, temp_dir)

                if images:
                    # 有图片，创建图片元素
                    for img_info in images:
                        elements.append(DocumentElement(
                            element_type='image',
                            content="",
                            metadata=img_info,
                            position=position
                        ))
                        position += 1

                # 添加段落文本
                if paragraph.text.strip():
                    elements.append(DocumentElement(
                        element_type='paragraph',
                        content=paragraph.text,
                        position=position
                    ))

        except Exception as e:
            logger.error(f"Failed to process paragraph: {e}")

        return elements

    def _process_table(
        self,
        ct_tbl: CT_Tbl,
        doc: Document,
        position: int
    ) -> Optional[DocumentElement]:
        """处理表格"""
        try:
            table = Table(ct_tbl, doc)

            # 提取表格数据
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            return DocumentElement(
                element_type='table',
                content="",
                metadata={
                    'rows': len(table_data),
                    'columns': len(table_data[0]) if table_data else 0,
                    'data': table_data
                },
                position=position
            )

        except Exception as e:
            logger.error(f"Failed to process table: {e}")
            return None

    def _extract_images_from_paragraph(
        self,
        paragraph: Paragraph,
        temp_dir: Optional[str]
    ) -> List[Dict[str, Any]]:
        """从段落中提取图片"""
        images = []
        image_index = 0

        if not temp_dir:
            return images

        try:
            # 遍历段落中的所有runs
            for run in paragraph.runs:
                # 查找图片关系ID
                for inline in run.element.xpath('.//a:blip'):
                    try:
                        # 获取图片嵌入关系ID
                        embed_id = inline.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                        )

                        if embed_id:
                            # 从文档的part关系中获取图片
                            document_part = paragraph._parent.part
                            image_part = document_part.related_parts[embed_id]

                            # 获取图片数据
                            image_data = image_part.blob
                            content_type = image_part.content_type
                            ext = content_type.split('/')[-1]  # e.g., 'png', 'jpeg'

                            # 保存图片到临时目录
                            import uuid
                            image_filename = f"docx_img_{uuid.uuid4().hex[:8]}_{image_index}.{ext}"
                            image_path = os.path.join(temp_dir, image_filename)

                            with open(image_path, 'wb') as f:
                                f.write(image_data)

                            images.append({
                                'type': 'embedded',
                                'temp_path': image_path,
                                'filename': image_filename,
                                'content_type': content_type
                            })

                            image_index += 1

                    except Exception as e:
                        logger.warning(f"Failed to extract individual image: {e}")
                        continue

        except Exception as e:
            logger.error(f"Failed to extract images from paragraph: {e}")

        return images

    def _extract_heading_level(self, style_name: str) -> int:
        """提取标题级别"""
        match = re.search(r'Heading\s*(\d+)', style_name, re.IGNORECASE)
        if match:
            return int(match.group(1))
        return 1

    def _organize_sections(self, elements: List[DocumentElement]) -> List[DocumentSection]:
        """组织为章节结构"""
        sections = []
        section_stack = []  # 用于跟踪当前章节层级

        for element in elements:
            if element.element_type == 'heading':
                # 创建新章节
                section = DocumentSection(
                    level=element.level,
                    title=element.content,
                    content=[],
                    metadata=element.metadata
                )

                # 找到合适的父章节
                while section_stack and section_stack[-1].level >= section.level:
                    section_stack.pop()

                if section_stack:
                    section_stack[-1].subsections.append(section)
                else:
                    sections.append(section)

                section_stack.append(section)

            elif element.element_type == 'paragraph':
                # 添加到当前章节
                if section_stack:
                    section_stack[-1].content.append(element.content)
                else:
                    # 没有章节，创建根章节
                    if not sections:
                        sections.append(DocumentSection(
                            level=0,
                            title="文档内容",
                            content=[element.content]
                        ))
                    else:
                        sections[0].content.append(element.content)

            elif element.element_type in ['image', 'table', 'formula', 'chart']:
                # 添加多媒体元素到当前章节
                if section_stack:
                    if element.element_type == 'image':
                        section_stack[-1].images.append(element.metadata)
                    elif element.element_type == 'table':
                        section_stack[-1].tables.append(element.metadata)
                    elif element.element_type == 'formula':
                        section_stack[-1].formulas.append(element.metadata)
                else:
                    # 添加到根章节
                    if not sections:
                        sections.append(DocumentSection(
                            level=0,
                            title="文档内容",
                            content=[]
                        ))
                    if element.element_type == 'image':
                        sections[0].images.append(element.metadata)
                    elif element.element_type == 'table':
                        sections[0].tables.append(element.metadata)
                    elif element.element_type == 'formula':
                        sections[0].formulas.append(element.metadata)

        return sections

    async def _analyze_multimedia_elements(
        self,
        elements: List[DocumentElement],
        temp_dir: Optional[str],
        enable_ocr: bool
    ):
        """分析多媒体元素"""
        if not self.multimodal_analyzer:
            return

        for element in elements:
            if element.element_type == 'image':
                image_path = element.metadata.get('temp_path')
                if image_path and os.path.exists(image_path):
                    # 使用多模态分析器分析图片
                    result = await self.multimodal_analyzer.analyze_image(
                        image_path,
                        image_type='auto'
                    )

                    # 更新元素内容和元数据
                    element.content = result.description
                    element.metadata.update({
                        'ocr_text': result.ocr_text,
                        'image_type': result.image_type,
                        'chart_info': result.chart_info,
                        'formula_info': result.formula_info,
                        'confidence': result.confidence
                    })

                    # 根据图片类型更新元素类型
                    if result.image_type == 'chart':
                        element.element_type = 'chart'
                    elif result.image_type == 'formula':
                        element.element_type = 'formula'

    def _generate_markdown(
        self,
        elements: List[DocumentElement],
        sections: List[DocumentSection]
    ) -> str:
        """生成Markdown输出"""
        markdown_parts = []

        # 方式1：按元素顺序生成
        for element in elements:
            if element.element_type == 'heading':
                markdown_parts.append(f"{'#' * element.level} {element.content}\n")
            elif element.element_type == 'paragraph':
                markdown_parts.append(f"{element.content}\n")
            elif element.element_type == 'image':
                img_desc = element.metadata.get('description', '图片')
                ocr_text = element.metadata.get('ocr_text', '')
                markdown_parts.append(f"\n**图片描述**: {img_desc}\n")
                if ocr_text:
                    markdown_parts.append(f"*图片中的文字*: {ocr_text}\n")
            elif element.element_type == 'chart':
                chart_desc = element.content
                chart_info = element.metadata.get('chart_info', {})
                markdown_parts.append(f"\n**图表**:\n\n{chart_desc}\n\n")
                if chart_info.get('x_axis'):
                    markdown_parts.append(f"- *横坐标*: {chart_info['x_axis']}\n")
                if chart_info.get('y_axis'):
                    markdown_parts.append(f"- *纵坐标*: {chart_info['y_axis']}\n")
                if chart_info.get('trend'):
                    markdown_parts.append(f"- *趋势*: {chart_info['trend']}\n")
            elif element.element_type == 'formula':
                formula_desc = element.content
                markdown_parts.append(f"\n**公式**: {formula_desc}\n")
            elif element.element_type == 'table':
                table_data = element.metadata.get('data', [])
                markdown_parts.append(f"\n{self._format_table_as_markdown(table_data)}\n")

        return "\n".join(markdown_parts)

    def _format_table_as_markdown(self, table_data: List[List[str]]) -> str:
        """将表格格式化为Markdown"""
        if not table_data:
            return ""

        # 确保第一行不是None
        first_row = table_data[0] or []
        if not first_row:
            return ""

        # 表头
        header = "| " + " | ".join(str(cell) if cell is not None else "" for cell in first_row) + " |"
        separator = "| " + " | ".join(["---"] * len(first_row)) + " |"

        # 表体
        rows = []
        for row in table_data[1:]:
            # 确保row不是None
            safe_row = row if row is not None else []
            rows.append("| " + " | ".join(str(cell) if cell is not None else "" for cell in safe_row) + " |")

        return "\n".join([header, separator] + rows)

    def _generate_json_output(
        self,
        elements: List[DocumentElement],
        sections: List[DocumentSection]
    ) -> str:
        """生成JSON输出"""
        import json

        output = {
            'elements': [
                {
                    'type': el.element_type,
                    'content': el.content,
                    'level': el.level,
                    'metadata': el.metadata,
                    'position': el.position
                }
                for el in elements
            ],
            'sections': self._sections_to_dict(sections)
        }

        return json.dumps(output, ensure_ascii=False, indent=2)

    def _sections_to_dict(self, sections: List[DocumentSection]) -> List[Dict]:
        """将章节列表转换为字典"""
        return [
            {
                'level': sec.level,
                'title': sec.title,
                'content': sec.content,
                'subsections': self._sections_to_dict(sec.subsections),
                'images': sec.images,
                'tables': sec.tables,
                'formulas': sec.formulas,
                'metadata': sec.metadata
            }
            for sec in sections
        ]

    def _build_metadata(
        self,
        file_path: str,
        elements: List[DocumentElement],
        sections: List[DocumentSection]
    ) -> Dict[str, Any]:
        """构建元数据"""
        file_path_obj = Path(file_path)

        # 统计各类元素数量
        element_counts = {}
        for element in elements:
            element_counts[element.element_type] = element_counts.get(element.element_type, 0) + 1

        metadata = {
            'file_type': 'docx',
            'parser_version': '2.0',
            'multimodal_enabled': self.enable_multimodal,
            'ocr_enabled': self.enable_ocr,
            'total_elements': len(elements),
            'element_counts': element_counts,
            'total_sections': len(sections),
            'headings': element_counts.get('heading', 0),
            'paragraphs': element_counts.get('paragraph', 0),
            'images': element_counts.get('image', 0),
            'charts': element_counts.get('chart', 0),
            'tables': element_counts.get('table', 0),
            'formulas': element_counts.get('formula', 0),
            'file_size': file_path_obj.stat().st_size if file_path_obj.exists() else 0
        }

        # 添加文档结构信息
        if sections:
            metadata['outline'] = self._generate_outline(sections)

        return metadata

    def _generate_outline(self, sections: List[DocumentSection], level: int = 0) -> List[Dict]:
        """生成文档大纲"""
        outline = []
        for section in sections:
            outline.append({
                'level': section.level,
                'title': section.title,
                'subsections': self._generate_outline(section.subsections, level + 1)
            })
        return outline
